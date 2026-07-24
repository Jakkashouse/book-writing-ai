# -*- coding: utf-8 -*-
"""함명자 작가님 출판권 설정 계약서 — 합의 사항 수정 조항 워드 생성기."""

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


OUT_PATH = r"C:\Users\JUN\Downloads\출판권 설정 계약서_함명자_수정조항.docx"


def set_korean_font(run, name="맑은 고딕", size=10, bold=False, color=None):
    run.font.name = name
    run.font.size = Pt(size)
    run.font.bold = bold
    if color is not None:
        run.font.color.rgb = color
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
        rPr.append(rFonts)
    rFonts.set(qn("w:eastAsia"), name)
    rFonts.set(qn("w:ascii"), name)
    rFonts.set(qn("w:hAnsi"), name)


def add_title(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(text)
    set_korean_font(r, size=18, bold=True)
    p.paragraph_format.space_after = Pt(6)


def add_subtitle(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(text)
    set_korean_font(r, size=11, bold=False, color=RGBColor(0x55, 0x55, 0x55))
    p.paragraph_format.space_after = Pt(12)


def add_h2(doc, text):
    p = doc.add_paragraph()
    r = p.add_run(text)
    set_korean_font(r, size=13, bold=True, color=RGBColor(0x1F, 0x3A, 0x68))
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after = Pt(4)


def add_label(doc, text, color=RGBColor(0x88, 0x44, 0x00)):
    p = doc.add_paragraph()
    r = p.add_run(text)
    set_korean_font(r, size=10, bold=True, color=color)
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(2)


def add_body(doc, text, indent_cm=0.3, italic=False, color=None):
    p = doc.add_paragraph()
    r = p.add_run(text)
    set_korean_font(r, size=10, bold=False, color=color)
    r.italic = italic
    p.paragraph_format.left_indent = Cm(indent_cm)
    p.paragraph_format.space_after = Pt(2)


def add_note(doc, text):
    p = doc.add_paragraph()
    r = p.add_run(text)
    set_korean_font(r, size=9.5, bold=False, color=RGBColor(0x55, 0x55, 0x55))
    r.italic = True
    p.paragraph_format.left_indent = Cm(0.3)
    p.paragraph_format.space_after = Pt(8)


def add_hr(doc):
    p = doc.add_paragraph()
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "888888")
    pBdr.append(bottom)
    pPr.append(pBdr)


doc = Document()
section = doc.sections[0]
section.top_margin = Cm(2.0)
section.bottom_margin = Cm(2.0)
section.left_margin = Cm(2.2)
section.right_margin = Cm(2.2)

add_title(doc, "출판권 설정 계약서 — 합의 사항 수정 조항")
add_subtitle(doc, "도서『사람안의 빛』· 저자 함명자 · 출판사 작가의집 · 2026년 4월 13일자 계약 보완")

# 개요
add_h2(doc, "■ 본 문서의 사용 방법")
add_body(
    doc,
    "본 문서는 2026년 4월 13일자 「출판권 및 전송권 설정 계약서」(저자 함명자, 출판사 작가의집)의 "
    "조항 중 양 당사자가 보완·수정에 합의한 사항을 정리한 것입니다. 아래 조항은 원 계약서의 "
    "해당 조항을 대체하거나 신설되며, 본 수정 조항과 원 계약서 사이에 충돌이 있는 경우 본 "
    "수정 조항이 우선합니다.",
)
add_body(
    doc,
    "양 당사자는 본 문서에 합의함으로써 원 계약서 제3조·제6조·제11조·제14조·제15조·제17조의 "
    "개정에 동의하며, 원 계약서 말미의 서명·날인란이 본 문서에도 동일하게 효력을 갖는 것으로 본다.",
)

# ─────────────────────────────────────────
add_h2(doc, "1. 제3조 (출판권의 설정) — ③항 신설")

add_label(doc, "[기존]")
add_body(doc, "① 저작권자는 출판사에게 위 저작물에 대한 출판권을 설정한다.")
add_body(
    doc,
    "② 제1항의 규정에 따라 출판사는 위 저작물을 원작 그대로 출판할 수 있는 독점적이고도 "
    "배타적인 권리를 가진다.",
)

add_label(doc, "[수정]", color=RGBColor(0x00, 0x66, 0x00))
add_body(doc, "① 저작권자는 출판사에게 위 저작물에 대한 출판권을 설정한다.")
add_body(
    doc,
    "② 제1항의 규정에 따라 출판사는 위 저작물을 원작 그대로 출판할 수 있는 독점적이고도 "
    "배타적인 권리를 가진다.",
)
add_body(
    doc,
    "③ 제2항에도 불구하고, 저작권자가 자신의 강연·강의·교육·세미나 및 본인의 SNS·블로그 "
    "등 개인 채널을 통한 통상적인 교육 활동과 홍보 활동의 범위 안에서 위 저작물의 본문 "
    "일부를 직접 인용·재사용하는 것은 출판권을 침해하지 아니하는 것으로 본다. 다만, 위 "
    "저작물의 동일성을 해치거나 출판물 전체를 대체할 정도의 분량을 재사용하여서는 아니 된다.",
)
add_note(doc, "근거: 작가님 1차 메일 1번 요청 → 대표님 1차 답변 1번 수락.")

# ─────────────────────────────────────────
add_h2(doc, "2. 제6조 (출판권의 존속기간 등) — ①·③항 수정")

add_label(doc, "[기존]")
add_body(
    doc,
    "① 출판사가 보유하는 위 저작물의 출판권은 계약일로부터 초판 1쇄 발행일까지, 그리고 "
    "초판 1쇄 발행일로부터 5년까지 효력을 가진다.",
)
add_body(
    doc,
    "③ 제2항에 따른 종료 통보가 없는 경우에 이 계약은 동일한 조건으로 5년까지 자동 연장되며, "
    "이 경우 출판사는 자동 연장 이전까지의 저작권사용료를 정산하여야 한다.",
)

add_label(doc, "[수정]", color=RGBColor(0x00, 0x66, 0x00))
add_body(
    doc,
    "① 출판사가 보유하는 위 저작물의 출판권은 계약일로부터 초판 1쇄 발행일까지, 그리고 "
    "초판 1쇄 발행일로부터 4년까지 효력을 가진다.",
)
add_body(
    doc,
    "③ 제2항에 따른 종료 통보가 없는 경우라도 이 계약은 자동으로 연장되지 아니하며, "
    "계약기간 만료 후 계약을 연장하고자 할 때에는 저작권자와 출판사가 서면으로 "
    "합의하여야 한다. 이 경우 인세율을 포함한 계약 조건은 연장 시점의 상황에 맞추어 "
    "양 당사자가 재협의하여 정한다. 출판사는 자동 연장 이전까지의 저작권사용료를 정산한다.",
)
add_note(
    doc,
    "근거: 작가님 1차 메일 2번 + 2차 메일 1번(기간 4년 재제안) → 대표님 자동연장 삭제 동의, "
    "기간은 4년으로 최종 조정.",
)

# ─────────────────────────────────────────
add_h2(doc, "3. 제11조 (저작물의 수정증감 및 비용 부담) — ④항 수정")

add_label(doc, "[기존]")
add_body(
    doc,
    "④ 초판 1쇄 발행 이후 중쇄 또는 중판을 발행함에 있어 저작자의 요청에 따른 수정, 증감 "
    "등에 의하여 통상의 제작비를 현저히 초과하는 경우 그 초과금액에 대한 저작권자의 "
    "부담액은 저작권자와 출판사가 협의하여 정한다. 이때 통상의 제작비는 초판 1쇄 발행 "
    "비용을 기준으로 산정한다.",
)

add_label(doc, "[수정]", color=RGBColor(0x00, 0x66, 0x00))
add_body(
    doc,
    "④ 초판 1쇄에서 발견된 오탈자 등 통상의 교정 사항은 출판사가 비용 부담 없이 재쇄 시 "
    "반영한다. 다만, 초판 1쇄 발행 이후 중쇄 또는 중판을 발행함에 있어 저작자의 요청에 "
    "따른 수정·증감 등에 의하여 통상의 제작비를 현저히 초과하는 경우, 그 초과금액에 "
    "대한 저작권자의 부담은 사전에 저작권자와 출판사가 서면으로 합의한 경우에 한하여 "
    "발생한다. 이때 통상의 제작비는 초판 1쇄 발행 비용을 기준으로 산정한다.",
)
add_note(doc, "근거: 작가님 1차 메일 3번 → 대표님 1차 답변 3번 전면 수락.")

# ─────────────────────────────────────────
add_h2(doc, "4. 제14조 (계속 출판의 의무 및 해지) — ②항 신설")

add_label(doc, "[기존]")
add_body(
    doc,
    "제14조 (계속 출판의 의무) 출판사는 이 계약기간 중 위 저작물을 계속 출판하여야 한다. "
    "다만, 6개월 동안 판매량이 100부 이하가 될 경우, 저작권자와 출판사가 합의하여 이 "
    "계약을 해지할 수 있다.",
)

add_label(doc, "[수정]", color=RGBColor(0x00, 0x66, 0x00))
add_body(
    doc,
    "① 출판사는 이 계약기간 중 위 저작물을 계속 출판하여야 한다. 다만, 6개월 동안 "
    "판매량이 100부 이하가 될 경우, 저작권자와 출판사가 합의하여 이 계약을 해지할 수 있다.",
)
add_body(
    doc,
    "② 제1항의 판매 부진을 사유로 저작권자가 출판사에 서면으로 계약 해지를 통보한 후 "
    "3개월 이내에 출판사가 재쇄 발행 등 위 저작물의 재판매를 위한 후속 조치를 취하지 "
    "아니하는 경우, 별도의 합의 없이 그 기간 만료일에 이 계약은 종료된 것으로 본다. "
    "이때 출판사는 제24조에 따른 재고품 배포 기간 동안 이를 처리할 수 있으며, 해당 "
    "기간에 대한 저작권사용료는 제15조 제1항에 따라 지급한다.",
)
add_note(
    doc,
    "근거: 작가님 2차 메일 1번 보완안(해지 통보 후 3개월 내 후속 조치 없으면 자동 종료). "
    "※ 대표님 최종 메시지의 '품절 우려' 확인 답변은 아직 회신 전이므로, 본 조항은 작가님 "
    "보완안을 그대로 반영한 안입니다. 작가님 회신 내용에 따라 추후 ②항 문구를 조정할 수 "
    "있습니다.",
)

# ─────────────────────────────────────────
add_h2(doc, "5. 제15조 (저작권사용료 등) — ①·③항 수정, ④항 신설")

add_label(doc, "[기존]")
add_body(
    doc,
    "① 출판사는 아래와 같이 저작권자에게 정가의 일정 비율에 해당하는 금액에 일정 부수"
    "(발행부수 또는 판매부수)를 곱한 금액을 지정 계좌를 통하여 저작권사용료로 지급한다. … "
    "인세는 1쇄부터 순출고부수 × 도서 정가의 10%를 지불한다. 순출고부수는 출고부수에서 "
    "반품 및 증정부수를 제외한 부수로 정산한다.",
)
add_body(
    doc,
    "③ (면제부수) 저작권자는 납본, 증정, 신간 안내, 서평, 홍보 등을 위하여 제공되는 부수에 "
    "대하여는 저작권사용료를 면제한다. 다만, 그 부수는 매쇄 당 10%를 초과할 수 없으며, "
    "출판사는 자세한 내역을 저작권자에게 알려주어야 한다.",
)

add_label(doc, "[수정]", color=RGBColor(0x00, 0x66, 0x00))
add_body(
    doc,
    "① 출판사는 아래와 같이 저작권자에게 저작권사용료를 지정 계좌를 통하여 지급한다. "
    "인세는 1쇄부터 순출고부수 × 도서 정가의 10%로 산정한다. 순출고부수는 출고부수에서 "
    "반품 및 증정부수를 제외한 부수로 한다.",
)
add_body(
    doc,
    "③ (면제부수) 저작권자는 납본, 증정, 신간 안내, 서평, 홍보 등을 위하여 제공되는 부수에 "
    "대하여는 저작권사용료를 면제한다. 다만, 그 부수는 매쇄 당 5%를 초과할 수 없으며, "
    "출판사는 자세한 내역을 저작권자에게 알려주어야 한다.",
)
add_body(
    doc,
    "④ (정산 리포트) 출판사는 매 정산기(6월·12월)마다 저작권사용료를 지급할 때 다음 각 "
    "호의 내용을 포함한 정산 리포트를 저작권자에게 서면(이메일 포함)으로 함께 제공한다.",
)
add_body(doc, "    1. 해당 정산기 동안의 발행부수·출고부수·반품부수·증정부수의 내역", indent_cm=0.6)
add_body(doc, "    2. 누적 발행부수 및 재고 현황", indent_cm=0.6)
add_body(doc, "    3. 산정 근거(도서 정가, 적용 인세율, 면제부수 사용 내역)", indent_cm=0.6)
add_body(doc, "    4. 해당 기간 저작권사용료 산정 금액 및 송금 내역", indent_cm=0.6)
add_note(
    doc,
    "근거: 작가님 1차 메일 5번(면제부수 10% → 5%) + 2차 메일 2번(정산 리포트 명문화) → "
    "대표님 모두 수락. 인세율 인상 기준(누적 3,000부 시 12%)은 대표님이 '추후 협의'로 "
    "유지하기로 하여 본 수정안에서는 제외함.",
)

# ─────────────────────────────────────────
add_h2(doc, "6. 제17조 (2차적저작물작성권 등) — 전면 개정")

add_label(doc, "[기존]")
add_body(
    doc,
    "① 이 계약기간 중에 위 저작물의 전자책(e-book) 및 오디오북 제작·배포·전송에 관한 "
    "권리는 출판사에게 있으며, 출판사는 이를 위하여 필요한 편집, 변환, 녹음, 제작 등의 "
    "행위를 할 수 있다.",
)
add_body(doc, "② 전자책 및 오디오북의 수익 배분은 다음과 같다. (전자책 50%, 오디오북 25% …)")
add_body(
    doc,
    "③ 전자책 및 오디오북을 제외한 기타 2차적저작물(번역, 각색, 편곡, 영상화 등)의 "
    "이용에 관하여는 다음과 같다. (저작권자에게 모든 권리 …)",
)
add_body(
    doc,
    "④ 제③항의 기타 2차적저작물 사용으로 발생한 수익은 순수익을 기준으로 '갑'과 '을'이 "
    "5:5로 배분한다.",
)
add_body(
    doc,
    "⑤ 출판사에 저작권대리중개업 자격이 있는 경우 저작권자는 기타 2차적 및 부차적 "
    "이용허락에 관한 사항의 전부 또는 일부를 출판사에게 위임할 수 있다.",
)

add_label(doc, "[수정]", color=RGBColor(0x00, 0x66, 0x00))
add_body(
    doc,
    "① 본 계약의 출판권은 종이책(인쇄·도화) 발행에 한정한다. 위 저작물의 전자책(e-book) 및 "
    "오디오북의 제작·배포·전송에 관한 권리는 본 계약의 대상에서 제외하며, 그 권리는 "
    "저작권자에게 유보된다. 전자책 및 오디오북의 제작·발행이 필요한 경우에는, 그 시점에 "
    "저작권자와 출판사가 별도의 서면 계약을 체결하여 수익 배분율을 포함한 구체적 조건을 "
    "정한다.",
)
add_body(
    doc,
    "② 전자책 및 오디오북을 제외한 기타 2차적저작물(번역, 각색, 편곡, 영상화 등)의 이용에 "
    "관하여는 다음과 같다.",
)
add_body(doc, "    1. 그에 관한 이용허락 등 모든 권리는 저작권자에게 있다.", indent_cm=0.6)
add_body(
    doc,
    "    2. 출판사에 먼저 요청이 오는 경우 출판사는 그 사실을 저작권자에게 즉시 알려주어야 한다.",
    indent_cm=0.6,
)
add_body(
    doc,
    "③ 제②항의 기타 2차적저작물 사용으로 발생한 수익은 순수익을 기준으로 '갑'과 '을'이 "
    "5:5로 배분하고, 해당기간 출판권 설정대가를 지급할 때 함께 정산해 지급한다.",
)
add_body(
    doc,
    "④ 제②항에도 불구하고 출판사에 저작권법에 따른 저작권대리중개업 자격이 있는 경우, "
    "저작권자는 기타 2차적 및 부차적 이용허락에 관한 사항의 전부 또는 일부를 출판사에게 "
    "위임할 수 있다.",
)
add_note(
    doc,
    "근거: 작가님 1차 메일 6번(전자책·오디오북은 별도 계약) → 대표님 1차 답변 6번 "
    "'해당 조항 삭제' 수락. 기존 ②항(전자책 50% / 오디오북 25% 배분) 전면 삭제. "
    "기존 ③·④·⑤항은 ②·③·④항으로 번호 재정렬.",
)

# ─────────────────────────────────────────
add_h2(doc, "■ 합의 요약표")
table = doc.add_table(rows=1, cols=3)
table.style = "Light Grid Accent 1"
hdr = table.rows[0].cells
for cell, text in zip(hdr, ["조항", "쟁점", "최종 합의"]):
    cell.text = ""
    p = cell.paragraphs[0]
    r = p.add_run(text)
    set_korean_font(r, size=10, bold=True)

rows = [
    ("제3조", "저자의 콘텐츠 활용권", "강연·SNS 통상적 활동 범위 내 일부 인용 허용 (③항 신설)"),
    ("제6조", "기간 / 자동연장", "5년 → 4년 / 자동연장 삭제, 서면 합의 시에만 연장"),
    ("제11조", "추가 비용 부담", "오탈자 재쇄 무비용 반영 + 추가비용은 사전 서면 합의 시에만"),
    ("제14조", "해지권 보완", "해지 통보 후 3개월 내 후속 조치 없으면 자동 종료"),
    ("제15조", "면제부수 / 정산 리포트", "면제부수 10% → 5%, 매 정산기 4개 항목 리포트 명문화"),
    ("제15조", "인세율 인상 기준", "본 수정안에서는 제외 (추후 성과 보며 협의)"),
    ("제17조", "전자책·오디오북", "본 계약에서 제외, 발행 시 별도 계약 체결"),
]
for row_data in rows:
    cells = table.add_row().cells
    for cell, text in zip(cells, row_data):
        cell.text = ""
        p = cell.paragraphs[0]
        r = p.add_run(text)
        set_korean_font(r, size=9.5)

# ─────────────────────────────────────────
add_h2(doc, "■ 서명")
add_body(doc, "위와 같이 양 당사자는 본 수정 조항에 합의하며, 본 문서는 원 계약서와 동일한 효력을 갖는다.")
add_body(doc, " ")
add_body(doc, "                              2026년      월      일")
add_body(doc, " ")
add_body(doc, "저작권자 :  함  명  자  (인)")
add_body(doc, "출판사   :  작가의집 대표 황준연  (인)")

doc.save(OUT_PATH)
print(f"saved: {OUT_PATH}")
