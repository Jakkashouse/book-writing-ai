"""
업로드 파일에서 텍스트를 뽑아내는 유틸 — .md/.txt/.docx/.pdf 지원.
.hwp/.hwpx는 지원하지 않음 (업로드 시 안내 필요).
"""
from __future__ import annotations

import io
from typing import IO

SUPPORTED_EXTS = {".md", ".txt", ".docx", ".pdf"}


def extract_text(filename: str, data: bytes) -> str:
    """파일 이름과 raw bytes를 받아 텍스트 반환.

    실패해도 예외를 올리지 않고 빈 문자열을 반환 (호출부에서 len 체크).
    """
    lower = filename.lower()

    if lower.endswith((".md", ".txt")):
        for enc in ("utf-8", "utf-8-sig", "cp949", "euc-kr"):
            try:
                return data.decode(enc)
            except UnicodeDecodeError:
                continue
        return ""

    if lower.endswith(".docx"):
        try:
            from docx import Document
        except ImportError:
            return ""
        try:
            doc = Document(io.BytesIO(data))
            parts = [p.text for p in doc.paragraphs if p.text]
            # 표 안의 텍스트도 수집
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        if cell.text:
                            parts.append(cell.text)
            return "\n".join(parts)
        except Exception:
            return ""

    if lower.endswith(".pdf"):
        try:
            import pdfplumber
        except ImportError:
            return ""
        try:
            with pdfplumber.open(io.BytesIO(data)) as pdf:
                return "\n".join((page.extract_text() or "") for page in pdf.pages)
        except Exception:
            return ""

    return ""


def is_supported(filename: str) -> bool:
    lower = filename.lower()
    return any(lower.endswith(ext) for ext in SUPPORTED_EXTS)
